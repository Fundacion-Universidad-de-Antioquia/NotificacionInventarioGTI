# -*- encoding: utf-8 -*-

#! /usr/bin/env python
#created by @ceapalaciosal
#under code Creative Commons

import docx
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH 
#import sys; sys.maxunicode
from funciones import *
from datetime import datetime
import locale
locale.setlocale(locale.LC_ALL,'es_CO.UTF-8')

def word(datos):

    id = datos.keys()

    for i in id:

        Nombre = (datos[i]['Nombre'].decode('iso-8859-1').encode('utf8')).upper()
        Apellidos = (datos[i]['Apellidos'].decode('iso-8859-1').encode('utf8')).upper()
        Cargo = datos[i]['Cargo']
        fecha = datetime.today().strftime('%d de %B del %Y')

        archive = '..' + '/' + 'archives' + '/' + 'Word' + '/' + str(i) + ".docx"
        doc = docx.Document()
        # doc.styles['Normal'].font.name = 'Arial'
        # doc.font.size = Pt(12)

        '''Apply style'''
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = docx.shared.Pt(8)

        #Logo de arriba
        doc.save(archive)
        header = doc.sections[0].header
        paragraph = header.paragraphs[0]
        logo_run = paragraph.add_run()
        logo_run.add_picture("img/Logo_arriba.png",width=Inches(6), height=Inches(1)) #width=Inches(7.5)

        #Footer insert
        Footer = doc.sections[0].footer
        paragraph = Footer.paragraphs[0]
        logo_run = paragraph.add_run()
        logo_run.add_picture("img/pie.jpg",width=Inches(6), height=Inches(1)) #width=Inches(5)

        texto = 'Medelín, '.decode("utf-8") + fecha

        para = doc.add_paragraph()
        para.add_run(texto)
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        texto = 'Señor(a):'.decode("utf-8")
        para = doc.add_paragraph()
        para.add_run(texto)
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        texto = ('\n' + Nombre + ' ' + Apellidos).decode("utf-8")
        para.add_run(texto).bold = True
        texto = '\nCargo: ' + Cargo.decode("utf-8")
        para.add_run(texto)
        para = doc.add_paragraph()
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        texto = 'Asunto: Notificación de elementos a cargo según inventario'.decode("utf-8")
        para = doc.add_paragraph()
        para.add_run(texto)
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        texto = 'Cordial saludo, la presente es para notificarle que, según el inventario de la Fundación Universidad de Antioquia, usted tiene bajo su responsabilidad los siguientes elementos:'.decode("utf-8")
        para = doc.add_paragraph()
        para.add_run(texto)
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Crea 5 filas y 7 columnas
        if  'N/A' == datos[i]['Equipos']['PC'] or {} == datos[i]['Equipos']['PC']:
            pass
        else:
            #Se escribe la tabla de equipos de computo
            texto = 'Equipo(s) de Computo'.decode("utf-8")
            para = doc.add_paragraph()
            para.add_run(texto).bold = True
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

            Equipos = datos[i]['Equipos']['PC']
            table = doc.add_table(rows=len(Equipos)+1, cols=6, style='Light Grid Accent 1')
            table.alignment=WD_ALIGN_PARAGRAPH.CENTER
            #Modificar el contenido de la celda de la segunda fila y la tercera columna a China
            table.cell(0,0).text = 'Número Inventario'.decode("utf-8")
            table.cell(0,1).text = 'Tipo'
            table.cell(0,2).text = 'Fabricante'
            table.cell(0,3).text = 'Serial'
            table.cell(0,4).text = 'Modelo'.decode("utf-8")
            table.cell(0,5).text = 'Estado'.decode("utf-8")

            #Añadir una nueva fila en la parte inferior de la tabla
            row = 1
            for NInventario in Equipos:     
                table.cell (row,0) .text = str(NInventario)
                table.cell (row,1) .text = str(Equipos[NInventario]['Tipo']).decode("utf-8")
                table.cell (row,2) .text = Equipos[NInventario]['Fabricante'].decode("utf-8")
                table.cell (row,3) .text = str(Equipos[NInventario]['Serie'])
                table.cell (row,4) .text = str(Equipos[NInventario]['Modelo']).decode("utf-8")
                table.cell (row,5) .text = Equipos[NInventario]['Estado'].decode("utf-8")
                row += 1
        
        
        if  'N/A' == datos[i]['Equipos']['CEL'] or {} == datos[i]['Equipos']['CEL']:
            pass
        else: 
            texto = '\nEquipo(s) Celular'.decode("utf-8")
            para = doc.add_paragraph()
            para.add_run(texto).bold = True
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

            Equipos = datos[i]['Equipos']['CEL']

            # Crea 5 filas y 7 columnas
            table = doc.add_table(rows=len(Equipos)+1, cols=4, style='Light Grid Accent 1')
            table.alignment=WD_ALIGN_PARAGRAPH.CENTER
            #Modificar el contenido de la celda de la segunda fila y la tercera columna a China
            table.cell(0,0).text = 'Número Inventario'.decode("utf-8")
            table.cell(0,1).text = 'Marca'
            table.cell(0,2).text = 'Serial'
            table.cell(0,3).text = 'Estado'

            #Añadir una nueva fila en la parte inferior de la tabla
            row = 1
            for NInventario in Equipos:      
                table.cell (row,0) .text = str(NInventario)
                table.cell (row,1) .text = str(Equipos[NInventario]['Fabricante']).decode("utf-8")
                table.cell (row,2) .text = str(Equipos[NInventario]['Serie']).decode("utf-8")                    
                table.cell (row,3) .text = str(Equipos[NInventario]['Estado']).decode("utf-8")
                row += 1

        if  'N/A' == datos[i]['Equipos']['IP'] or {} == datos[i]['Equipos']['IP']:
            pass
        else: 
            texto = '\nTelefonos IP'.decode("utf-8")
            para = doc.add_paragraph()
            para.add_run(texto).bold = True
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

            Telefonos = datos[i]['Equipos']['IP']

            # Crea 5 filas y 7 columnas
            table = doc.add_table(rows=len(Telefonos)+1, cols=4, style='Light Grid Accent 1')
            table.alignment=WD_ALIGN_PARAGRAPH.CENTER
            #Modificar el contenido de la celda de la segunda fila y la tercera columna a China
            table.cell(0,0).text = 'Número Inventario'.decode("utf-8")
            table.cell(0,1).text = 'Marca'
            table.cell(0,2).text = 'Serial'
            table.cell(0,3).text = 'Estado'

            #Añadir una nueva fila en la parte inferior de la tabla
            row = 1
            for NInventario in Telefonos:          
                table.cell (row,0) .text = str(NInventario)
                table.cell (row,1) .text = str(Telefonos[NInventario]['Fabricante']).decode("utf-8")
                table.cell (row,2) .text = str(Telefonos[NInventario]['Serie']).decode("utf-8")                    
                table.cell (row,3) .text = str(Telefonos[NInventario]['Estado']).decode("utf-8")
                row += 1
        
        if  'N/A' == datos[i]['Equipos']['IMP'] or {} == datos[i]['Equipos']['IMP']:
            pass
        else: 
            texto = '\nImpresoras'.decode("utf-8")
            para = doc.add_paragraph()
            para.add_run(texto).bold = True
            para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            Impresoras = datos[i]['Equipos']['IMP']

            # Crea 5 filas y 7 columnas
            table = doc.add_table(rows=len(Impresoras)+1, cols=4, style='Light Grid Accent 1')
            table.alignment=WD_ALIGN_PARAGRAPH.CENTER
            #Modificar el contenido de la celda de la segunda fila y la tercera columna a China
            table.cell(0,0).text = 'Número Inventario'.decode("utf-8")
            table.cell(0,1).text = 'Marca'
            table.cell(0,2).text = 'Serial'
            table.cell(0,3).text = 'Estado'

            #Añadir una nueva fila en la parte inferior de la tabla
            row = 1
            for NInventario in Impresoras:          
                table.cell (row,0) .text = str(NInventario)
                table.cell (row,1) .text = str(Impresoras[NInventario]['Fabricante']).decode("utf-8")
                table.cell (row,2) .text = str(Impresoras[NInventario]['Serie']).decode("utf-8")                    
                table.cell (row,3) .text = str(Impresoras[NInventario]['Estado']).decode("utf-8")
                row += 1


        texto = '\n\nRecuerde que los elementos y sus componentes anteriormente listados, se encuentran a su cargo y por tanto usted es responsable por daños, perdidas u otros que no estén contemplados dentro del desgaste normal de los elementos, como se estipula en el reglamento interno de trabajo de la organización. Igualmente, si tiene alguna duda u objeción con alguno de los elementos listados puede informar al equipo de Gestión Tecnológica e Innovación para poder revisar su caso. '.decode("utf-8")
        para = doc.add_paragraph()
        para.add_run(texto)
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

       
        texto = '\nAtentamente,'.decode("utf-8")
        para = doc.add_paragraph()
        para.add_run(texto)
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        imagen = "img/Firma_Digital.png"
        para = doc.add_picture(imagen,width=Inches(1), height=Inches(0.8))
        
        texto1 = 'CESAR AUGUSTO PALACIOS ALARCON \nGerente de Tecnología e Innovación \nFundación Universidad de Antioquia'.decode("utf-8")
        para = doc.add_paragraph()
        para.add_run(texto1).bold = True
        para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

        doc.save(archive)